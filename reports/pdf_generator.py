import os
from docx import Document
from io import BytesIO

def generar_pdf_informe(datos):
    """
    Genera un documento Word basado en plantilla utilizando python-docx,
    reemplazando etiquetas dinámicas, aplicando fuente Calibri y quitando negritas a los datos.
    """
    plantilla_path = "reports/plantilla_base.docx"
    
    if not os.path.exists(plantilla_path):
        plantilla_path = "plantilla_base.docx"
        
    doc = Document(plantilla_path)

    reemplazos = {
        "{{fecha}}": str(datos.get("fecha_inicio", "")),
        "{{hora_inicio}}": str(datos.get("hora_inicio", "")),
        "{{hora_fin}}": str(datos.get("hora_fin", "En Proceso")),
        "{{metodo}}": str(datos.get("metodo", "Óxido de Etileno")),
        "{{control_carga}}": str(datos.get("control_carga", "Conforme")),
        "{{esterilizador}}": str(datos.get("esterilizador", "")),
        "{{n_cic}}": str(datos.get("n_cic", "")),
        "{{tipo_carga}}": str(datos.get("tipo_carga", "Textil")),
        "{{presion}}": str(datos.get("presion", "")),
        "{{tiempo_exposicion}}": str(datos.get("tiempo_exposicion", "")),
        "{{temperatura}}": str(datos.get("temperatura", "")),
        "{{operador}}": str(datos.get("operador", "")),
        "{{estado}}": str(datos.get("estado", ""))
    }

    def procesar_parrafo(p):
        texto_completo = "".join([run.text for run in p.runs])
        modificado = False
        for clave, valor in reemplazos.items():
            if clave in texto_completo:
                texto_completo = texto_completo.replace(clave, str(valor))
                modificado = True
        
        if modificado and p.runs:
            p.runs[0].text = texto_completo
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.bold = False
            for run in p.runs[1:]:
                run.text = ""
        else:
            for run in p.runs:
                run.font.name = 'Calibri'

    for p in doc.paragraphs:
        procesar_parrafo(p)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    procesar_parrafo(p)

    items = datos.get("items", [])
    
    if items and len(doc.tables) > 0:
        tabla_items = None
        for tabla in doc.tables:
            if len(tabla.rows) > 1:
                tabla_items = tabla
                break
        
        if tabla_items and len(tabla_items.rows) >= 2:
            for idx, item in enumerate(items):
                if idx + 1 < len(tabla_items.rows):
                    row = tabla_items.rows[idx + 1]
                else:
                    row = tabla_items.add_row()
                
                try:
                    row.cells[0].text = str(item.get("nombre", ""))
                    row.cells[1].text = str(item.get("cantidad", 0))
                    row.cells[2].text = str(item.get("lote", ""))
                    
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.name = 'Calibri'
                                run.font.bold = False
                except Exception:
                    pass

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generar_orden_entrega(datos):
    """
    Genera la Orden de Entrega localizando la fila de marcadores y rellenando 
    exclusivamente las celdas únicas de la tabla para evitar desajustes por colspan.
    """
    plantilla_path = "reports/plantilla_orden_entrega.docx"
    
    if not os.path.exists(plantilla_path):
        plantilla_path = "plantilla_orden_entrega.docx"
        
    doc = Document(plantilla_path)

    reemplazos = {
        "{{fecha}}": str(datos.get("fecha", "")),
        "{{n_orden}}": str(datos.get("n_orden", "0000")),
        "{{destino}}": str(datos.get("destino", "AM MEDICAL")),
        "{{ciclo_esteril}}": str(datos.get("ciclo_esteril", "")),
        "{{resp_entrega}}": str(datos.get("resp_entrega", "Davinson Peña")),
        "{{resp_recepcion}}": str(datos.get("resp_recepcion", "Jorge Espejero"))
    }

    def procesar_parrafo_orden(p):
        texto_completo = "".join([run.text for run in p.runs])
        modificado = False
        for clave, valor in reemplazos.items():
            if clave in texto_completo:
                texto_completo = texto_completo.replace(clave, str(valor))
                modificado = True
        
        if modificado and p.runs:
            p.runs[0].text = texto_completo
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.bold = False
            for run in p.runs[1:]:
                run.text = ""
        else:
            for run in p.runs:
                run.font.name = 'Calibri'

    # Procesar párrafos generales y cabeceras
    for p in doc.paragraphs:
        procesar_parrafo_orden(p)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    procesar_parrafo_orden(p)

    # Inyección de ítems usando celdas únicas por fila
    items = datos.get("items", [])
    
    if items:
        tabla_items = None
        fila_plantilla_idx = -1
        
        for tabla in doc.tables:
            for i, fila in enumerate(tabla.rows):
                texto_fila = "".join([c.text for c in fila.cells])
                if "{{nombre_item}}" in texto_fila or "nombre_item" in texto_fila:
                    tabla_items = tabla
                    fila_plantilla_idx = i
                    break
            if tabla_items:
                break
        
        if tabla_items and fila_plantilla_idx != -1:
            for idx, item in enumerate(items):
                target_idx = fila_plantilla_idx + idx
                
                if target_idx < len(tabla_items.rows):
                    row = tabla_items.rows[target_idx]
                else:
                    row = tabla_items.add_row()
                
                # Blindaje para detenerse si alcanza las filas de firmas institucionales
                texto_row = "".join([c.text for c in row.cells])
                if "Entregó" in texto_row or "Recibió" in texto_row or "Gerencia" in texto_row or "Nombre:" in texto_row:
                    break
                
                try:
                    # Filtrar celdas repetidas por colspan para obtener las 4 columnas visuales exactas
                    unique_cells = []
                    for c in row.cells:
                        if not unique_cells or c != unique_cells[-1]:
                            unique_cells.append(c)
                    
                    if len(unique_cells) >= 4:
                        # Asignación estricta y limpia a cada columna visual:
                        # [0] Descripción, [1] Lote/Ciclo, [2] Cantidad, [3] Observaciones
                        unique_cells[0].text = str(item.get("nombre", ""))
                        unique_cells[1].text = str(item.get("lote", ""))
                        unique_cells[2].text = str(item.get("cantidad", 0))
                        unique_cells[3].text = str(item.get("obs", "Conforme"))
                        
                        # Aplicar fuente Calibri estándar sin negritas
                        for cell in unique_cells[:4]:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.name = 'Calibri'
                                    run.font.bold = False
                except Exception as e:
                    print(f"Error escribiendo ítem {idx}: {e}")

            # Limpiar filas sobrantes hacia abajo hasta llegar al bloque de firmas
            siguiente_idx = fila_plantilla_idx + len(items)
            while siguiente_idx < len(tabla_items.rows):
                row = tabla_items.rows[siguiente_idx]
                texto_row = "".join([c.text for c in row.cells])
                if "Entregó" in texto_row or "Recibió" in texto_row or "Gerencia" in texto_row or "Nombre:" in texto_row:
                    break
                try:
                    for cell in row.cells:
                        cell.text = ""
                except:
                    pass
                siguiente_idx += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer